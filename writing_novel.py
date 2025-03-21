import tkinter as tk
from tkinter import scrolledtext, messagebox, ttk, filedialog
import requests
import json
import threading
import re
import time
import os
import logging
from datetime import datetime
import pyperclip  # 用于复制到剪贴板
import markdown  # 用于转换Markdown为HTML
from tkhtmlview import HTMLScrolledText  # 用于显示HTML
import docx  # 用于创建Word文档
from docx.shared import Pt

# 设置日志记录
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("novel_app.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("NovelApp")

# 全局变量，用于存储已生成的内容
generated_content = ""
thinking_content = ""  # 用于存储思维推理内容
is_generating = False  # 控制生成过程的标志
update_timer = None  # 用于存储更新计时器的ID
model_name = 'huihui_ai/qwen2.5-1m-abliterated:14b'  # 模型名称 
#writing_model_name = 'glm4:latest'
writing_model_name = 'huihui_ai/qwen2.5-1m-abliterated:14b'
#evaluation_model_name = 'qwq:latest'  # 用于评估的模型
evaluation_model_name = 'huihui_ai/qwen2.5-1m-abliterated:14b'  # 用于评估的模型
is_auto_generating = False  # 控制自动生成的标志
is_evaluating = False  # 控制评估过程的标志
is_markdown_mode = False  # 控制是否显示为Markdown格式

# 判断写作目标是否完成的方法
def is_writing_complete(content, target_word_count):
    # 计算中文字符数（每个中文字符算一个字）
    chinese_chars = re.findall(r'[\u4e00-\u9fff]', content)
    # 计算英文单词数（假设英文单词由空格分隔）
    english_words = re.findall(r'\b[a-zA-Z]+\b', content)
    
    # 总字数 = 中文字符数 + 英文单词数
    total_words = len(chinese_chars) + len(english_words)
    
    # 判断是否达到目标字数
    return total_words >= target_word_count

# 生成写作提示的方法
def generate_user_prompt():
    try:
        api_url = 'http://localhost:11434/api/generate'
        prompt = '''你是一个创意写作专家，请生成一个有趣的小说写作要求。要求：
        1. 包含具体的故事背景、人物设定和情节方向
        2. 要有创意，不要太过俗套
        3. 字数在100-200之间
        4. 只输出写作要求本身，不要包含任何其他说明
        5. 每次生成的内容都要不一样，保持新颖性
        6. 小说风格以言情类为主。
        '''
        
        request_data = {
            "model": model_name,
            "prompt": prompt,
            "temperature": 0.9,  # 使用较高的温度以增加创意性
            "stream": False
        }
        
        response = requests.post(api_url, json=request_data, timeout=30)
        response.raise_for_status()
        
        # 获取生成的提示词
        result = response.json()
        new_prompt = result.get('response', '').strip()
        
        # 清空并更新提示词输入框
        prompt_entry.delete("1.0", tk.END)
        prompt_entry.insert("1.0", new_prompt)
        
        return True
    except Exception as e:
        update_status(f"生成写作要求时出错：{str(e)}")
        return False

# 保存生成的内容到文件
def save_content_to_file():
    if not generated_content.strip():
        return
        
    # 确保输出目录存在
    output_dir = "generated_novels"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 计算当前内容的字数
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', generated_content))
    english_words = len(re.findall(r'\b[a-zA-Z]+\b', generated_content))
    total_words = chinese_chars + english_words
    
    # 生成文件名：字数_时间戳.txt
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{total_words}字_{timestamp}.txt"
    filepath = os.path.join(output_dir, filename)
    
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(generated_content)
        update_status(f"内容已保存至：{filename}")
        
        # 如果是自动生成模式，则继续生成下一个故事
        if is_auto_generating:
            root.after(2000, lambda: continue_auto_generate())
            
    except Exception as e:
        update_status(f"保存文件时出错：{str(e)}")

# 继续自动生成的方法
def continue_auto_generate():
    if is_auto_generating and not is_generating:
        if generate_user_prompt():
            generate_text()

# 处理本次生成的内容
def process_text_chunk(text_chunk):
    """处理文本块，分离思维推理内容和正式小说内容"""
    global thinking_content
    
    # 修正可能错误的</thind>标签为正确的</think>
    text_chunk = text_chunk.replace("</thind>", "</think>")
    
    # 存储原始文本块，用于比较处理前后的变化
    original_chunk = text_chunk
    story_content = ""
    
    # 循环处理可能存在的多个思维标签
    while "<think>" in text_chunk:
        # 找到<think>的位置
        think_start = text_chunk.find("<think>")
        
        # 将<think>前的内容添加到故事内容
        story_content += text_chunk[:think_start]
        
        # 截取剩余部分
        text_chunk = text_chunk[think_start:]
        
        # 找到</think>的位置
        think_end = text_chunk.find("</think>")
        
        if think_end != -1:
            # 完整的思维内容
            think_end += 8  # 加上"</think>"的长度(8个字符)
            think_content = text_chunk[:think_end]
            
            # 更新思维内容显示区域
            thinking_text.insert(tk.END, think_content)
            thinking_text.see(tk.END)
            
            # 将思维内容添加到思维记录中
            thinking_content += think_content
            
            # 处理剩余文本
            text_chunk = text_chunk[think_end:]
        else:
            # <think>标签没有配对的</think>，将整段视为思维内容
            thinking_text.insert(tk.END, text_chunk)
            thinking_text.see(tk.END)
            thinking_content += text_chunk
            text_chunk = ""
            break
    
    # 添加最后一部分文本到故事内容
    story_content += text_chunk
    
    # 检查是否正在一个未闭合的思维内容中
    if ("<think>" in thinking_content and "</think>" not in thinking_content) or \
       (thinking_content.rfind("<think>") > thinking_content.rfind("</think>")):
        # 可能是思维内容的中间部分，无标签
        if original_chunk != story_content:  # 已经处理过标签的不重复处理
            return story_content
        # 假设这是思维内容的一部分
        thinking_text.insert(tk.END, original_chunk)
        thinking_text.see(tk.END)
        thinking_content += original_chunk
        return ""
    
    return story_content

# 生成文本的线程函数
def generate_text_thread():
    global generated_content, thinking_content, is_generating, update_timer
    
    try:
        # 获取用户输入的提示和目标字数
        user_prompt = prompt_entry.get("1.0", tk.END).strip()
        try:
            target_word_count = int(word_count_entry.get())
        except ValueError:
            messagebox.showerror("错误", "请输入有效的目标字数")
            is_generating = False
            return
        
        # 清空显示区域
        output_text.delete(1.0, tk.END)
        thinking_text.delete(1.0, tk.END)
        generated_content = ""
        thinking_content = ""
        
        # 循环生成文本，直到达到目标字数
        while not is_writing_complete(generated_content, target_word_count) and is_generating:
            # 构造提示词（包括已生成的内容）
            if generated_content:
                full_prompt = f'''
                你是一个小说作家，擅长写言情类小说，请根据以下用户写作要求和已经写作的内容，开始或者继续创作：
                ## 用户的写作要求：
                 {user_prompt}

                ## 已经写作的内容：
                {generated_content}

                ## 注意事项：
                1. 请根据用户写作要求和已经写作的内容，继续创作。
                2. 请保持故事的连贯性和逻辑性。
                3. 请保持故事的节奏感，不要出现过于冗长或重复的描述。
                4. 请保持故事的紧凑性，不要出现过于拖沓的情节。
                5. 请保持故事的新鲜感，不要出现过于俗套的情节。
                6. 请保持故事的合理性，不要出现过于夸张的情节。
                7. 请保使用中文写作；
                8. 你可以使用<think></think>标签来表示你的思考过程，这部分内容不会出现在最终故事中。
                '''
            else:
                full_prompt = f'''{user_prompt}

你可以使用<think></think>标签来表示你的思考过程，这部分内容不会出现在最终故事中。例如：
<think>我需要先确定故事的主角和背景，然后设计一个合理的情节发展。</think>
正式的故事内容...'''
            
            # 准备请求参数
            api_url = 'http://localhost:11434/api/generate'
            request_data = {
                "model": writing_model_name,
                "prompt": full_prompt,
                "max_tokens": 1000000,
                "temperature": 0.7,
                "stream": True
            }
            
            # 发送请求并获取流式响应
            response = requests.post(api_url, json=request_data, stream=True, timeout=30)
            response.raise_for_status()
            
            # 处理本次生成的内容
            new_content = ""
            
            # 处理流式响应
            for line in response.iter_lines():
                # 如果已停止生成，则跳出响应处理循环
                if not is_generating:
                    break
                    
                if line:
                    # 解析 JSON 数据
                    data = json.loads(line.decode('utf-8'))
                    # 提取 response 字段
                    text_chunk = data.get('response', '')
                    
                    # 处理文本块，分离思维推理和正式内容
                    story_chunk = process_text_chunk(text_chunk)
                    
                    # 更新显示区域（仅显示正式内容）
                    if story_chunk:
                        output_text.insert(tk.END, story_chunk)
                        output_text.see(tk.END)  # 自动滚动到最新内容
                        # 累积新内容
                        new_content += story_chunk
                    
                    root.update_idletasks()  # 更新UI
                    
                    # 检查是否完成（当 done 为 true 时）
                    if data.get('done', False):
                        break
            
            # 如果已停止生成，则退出主循环
            if not is_generating:
                update_status("生成已停止")
                break
                
            # 更新已生成的内容
            generated_content += new_content
            
            # 检查是否应该继续生成
            if not is_generating:
                update_status("生成已停止")
                break
                
            # 模拟文本处理的延迟，避免过快请求
            root.after(1000)
        
        if is_generating and is_writing_complete(generated_content, target_word_count):
            update_status(f"写作完成！共生成{len(generated_content)}字")
            save_content_to_file()  # 自动保存内容
            
        elif not is_generating:
            update_status(f"用户已停止生成。当前已生成{len(generated_content)}字")
            save_content_to_file()  # 自动保存内容
        
    except requests.exceptions.RequestException as e:
        output_text.insert(tk.END, f"Error: {str(e)}\n")
        update_status("生成过程出错")
    except Exception as e:
        output_text.insert(tk.END, f"Error: {str(e)}\n")
        update_status("生成过程出错")
    finally:
        is_generating = False

# 定期更新状态的函数
def periodic_status_update():
    global update_timer
    if is_generating:
        word_count = len(re.findall(r'[\u4e00-\u9fff]', generated_content)) + len(re.findall(r'\b[a-zA-Z]+\b', generated_content))
        update_status(f"正在生成中...当前已生成约{word_count}字")
        # 每500毫秒更新一次状态
        update_timer = root.after(500, periodic_status_update)
    else:
        # 如果不再生成，则停止定期更新
        if update_timer:
            root.after_cancel(update_timer)
            update_timer = None

# 调用模型并更新显示区域的函数
def generate_text():
    global generated_content, is_generating, update_timer
    
    if is_generating:
        messagebox.showinfo("提示", "正在生成中，请稍候...")
        return
    
    is_generating = True
    
    # 开始定期更新状态
    periodic_status_update()
    
    # 创建一个新线程来执行生成过程
    thread = threading.Thread(target=generate_text_thread)
    thread.daemon = True
    thread.start()

# 停止生成的函数
def stop_generation():
    global is_generating
    is_generating = False
    update_status("用户已停止生成")

# 自动生成的处理函数
def auto_generate():
    global is_auto_generating
    
    if is_generating:
        messagebox.showinfo("提示", "正在生成中，请稍候...")
        return
    
    is_auto_generating = True
    auto_generate_button.config(text="停止自动生成")
    
    # 开始第一轮生成
    if generate_user_prompt():
        generate_text()

# 停止自动生成
def stop_auto_generate():
    global is_auto_generating
    is_auto_generating = False
    auto_generate_button.config(text="自动生成")
    stop_generation()

# 切换自动生成状态
def toggle_auto_generate():
    if is_auto_generating:
        stop_auto_generate()
    else:
        auto_generate()

# 评估小说质量的函数
def evaluate_novel_quality():
    global is_evaluating
    
    if not generated_content.strip():
        logger.warning("尝试评估空内容")
        messagebox.showinfo("提示", "请先生成小说内容再进行评估")
        return
    
    if is_evaluating:
        logger.warning("评估已在进行中，用户尝试重复评估")
        messagebox.showinfo("提示", "正在评估中，请稍候...")
        return
    
    content_length = len(generated_content)
    logger.info(f"开始评估小说质量，内容长度: {content_length} 字符")
    
    is_evaluating = True
    update_status("正在评估小说质量...")
    
    # 检查Ollama服务可用性
    try:
        logger.info("检查Ollama服务可用性...")
        response = requests.get('http://localhost:11434/api/tags', timeout=5)
        response.raise_for_status()
        
        # 检查评估模型是否可用
        models = response.json().get("models", [])
        model_names = [model.get("name") for model in models]
        
        if evaluation_model_name not in model_names:
            logger.warning(f"评估模型 {evaluation_model_name} 不在可用模型列表中")
            logger.info(f"可用模型: {model_names}")
            messagebox.showwarning("模型不可用", f"评估模型 '{evaluation_model_name}' 似乎不可用。\n\n请确保该模型已下载，或在代码中更改评估模型名称。\n\n可用模型: {', '.join(model_names)}")
            # 继续评估，因为模型名称可能是正确的，只是API返回的格式不同
    except Exception as e:
        logger.error(f"检查Ollama服务时出错: {str(e)}")
        if not messagebox.askyesno("服务检查失败", f"无法连接到Ollama服务或检查模型: {str(e)}\n\n是否仍要继续评估?"):
            is_evaluating = False
            update_status("评估已取消")
            return
    
    # 创建一个新线程来执行评估过程
    thread = threading.Thread(target=evaluate_novel_thread)
    thread.daemon = True
    thread.start()
    logger.info("评估线程已启动")

# 评估小说质量的线程函数
def evaluate_novel_thread():
    global is_evaluating
    
    try:
        logger.info("开始小说质量评估过程")
        user_prompt = prompt_entry.get("1.0", tk.END).strip()
        
        # 构造评估提示词
        evaluation_prompt = f'''
        请对以下小说内容进行专业的质量评估，基于以下几个方面:
        1. 情节连贯性和合理性
        2. 人物刻画和发展
        3. 写作风格和语言表达
        4. 创意性和独特性
        5. 与用户写作要求的符合度
        
        ## 用户的写作要求:
        {user_prompt}
        
        ## 小说内容:
        {generated_content[:2000]}...（内容较长，此处截断）
        
        请给出详细评价，并提出具体的改进建议。评分标准为1-10分，请在每个方面打分，并给出总体评分。
        格式要求:
        - 总体评分: X/10
        - 情节评分: X/10
        - 人物评分: X/10
        - 语言评分: X/10
        - 创意评分: X/10
        - 需求符合度: X/10
        
        ## 详细点评:
        [详细说明优点和不足]
        
        ## 具体改进建议:
        [列出3-5条具体的改进建议]
        '''
        
        logger.info(f"评估提示词长度: {len(evaluation_prompt)} 字符")
        
        # 准备请求参数
        api_url = 'http://localhost:11434/api/generate'
        request_data = {
            "model": evaluation_model_name,
            "prompt": evaluation_prompt,
            "temperature": 0.3,  # 使用较低的温度以获得更客观的评估
            "stream": False,
            "options": {
                "num_ctx": 4096  # 增加上下文窗口大小
            }
        }
        
        logger.info(f"使用评估模型: {evaluation_model_name}")
        logger.info("发送评估请求...")
        
        # 记录请求开始时间
        start_time = time.time()
        
        # 发送请求并增加超时时间
        response = requests.post(api_url, json=request_data, timeout=600)
        
        # 记录请求结束时间和耗时
        end_time = time.time()
        elapsed_time = end_time - start_time
        logger.info(f"评估请求完成，耗时: {elapsed_time:.2f} 秒")
        
        response.raise_for_status()
        
        # 获取评估结果
        result = response.json()
        evaluation_result = result.get('response', '').strip()
        
        logger.info(f"成功获取评估结果，长度: {len(evaluation_result)} 字符")
        
        # 显示评估结果
        root.after(0, lambda: show_evaluation_result(evaluation_result))
        
        update_status("小说质量评估完成")
    except requests.exceptions.Timeout as e:
        error_message = f"评估请求超时: {str(e)}"
        logger.error(error_message)
        root.after(0, lambda: messagebox.showerror("超时错误", f"评估过程超时: {str(e)}\n\n请检查以下可能的原因:\n1. 模型 '{evaluation_model_name}' 是否已加载\n2. Ollama服务器是否正常运行\n3. 是否有其他进程占用了大量资源"))
        root.after(0, lambda: update_status("评估过程超时"))
    except requests.exceptions.ConnectionError as e:
        error_message = f"评估请求连接错误: {str(e)}"
        logger.error(error_message)
        root.after(0, lambda: messagebox.showerror("连接错误", f"无法连接到Ollama服务: {str(e)}\n\n请检查Ollama服务是否正在运行"))
        root.after(0, lambda: update_status("评估连接失败"))
    except Exception as e:
        error_message = f"评估过程出错: {str(e)}"
        logger.error(error_message)
        logger.exception("评估过程详细错误")
        root.after(0, lambda: messagebox.showerror("错误", f"评估过程出错: {str(e)}"))
        root.after(0, lambda: update_status("评估过程出错"))
    finally:
        is_evaluating = False
        logger.info("评估过程结束")

# 显示评估结果的函数
def show_evaluation_result(evaluation_result):
    # 创建评估结果窗口
    result_window = tk.Toplevel(root)
    result_window.title("小说质量评估结果")
    result_window.geometry("800x600")
    result_window.configure(bg='#f5f5f7')
    
    # 创建主容器
    result_container = tk.Frame(result_window, bg='#f5f5f7')
    result_container.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)
    
    # 创建标题
    tk.Label(
        result_container,
        text="小说质量评估报告",
        font=('Microsoft YaHei UI', 16, 'bold'),
        bg='#f5f5f7',
        fg='#333333'
    ).pack(pady=(0, 15))
    
    # 创建评估结果文本区域
    result_text = scrolledtext.ScrolledText(
        result_container,
        wrap=tk.WORD,
        width=80,
        height=20,
        font=('Microsoft YaHei UI', 11),
        bg='#ffffff',
        fg='#333333',
        padx=10,
        pady=10
    )
    result_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
    result_text.insert(tk.END, evaluation_result)
    result_text.config(state=tk.DISABLED)  # 设置为只读
    
    # 提取评分信息
    try:
        total_score = re.search(r'总体评分:\s*(\d+(?:\.\d+)?)/10', evaluation_result)
        total_score = float(total_score.group(1)) if total_score else 0
        
        # 根据评分设置不同的颜色
        score_label = tk.Label(
            result_container,
            text=f"总体评分: {total_score}/10",
            font=('Microsoft YaHei UI', 14, 'bold'),
            padx=15,
            pady=10,
            relief=tk.GROOVE
        )
        
        if total_score >= 8:
            score_label.config(bg='#e8f5e9', fg='#2e7d32')  # 绿色 - 优秀
        elif total_score >= 6:
            score_label.config(bg='#fff8e1', fg='#f57f17')  # 琥珀色 - 良好
        else:
            score_label.config(bg='#ffebee', fg='#c62828')  # 红色 - 需要改进
            
        score_label.pack(pady=(15, 0))
    except:
        pass
    
    # 创建按钮区域
    button_frame = tk.Frame(result_container, bg='#f5f5f7')
    button_frame.pack(pady=15)
    
    # 添加修改建议按钮
    tk.Button(
        button_frame,
        text="获取修改建议",
        command=lambda: generate_revision_suggestions(evaluation_result),
        font=('Microsoft YaHei UI', 10, 'bold'),
        bg='#007aff',
        fg='white',
        relief=tk.RAISED,
        bd=0,
        padx=20,
        pady=8,
        cursor="hand2"
    ).pack(side=tk.LEFT, padx=10)
    
    # 添加关闭按钮
    tk.Button(
        button_frame,
        text="关闭",
        command=result_window.destroy,
        font=('Microsoft YaHei UI', 10, 'bold'),
        bg='#8e8e93',
        fg='white',
        relief=tk.RAISED,
        bd=0,
        padx=20,
        pady=8,
        cursor="hand2"
    ).pack(side=tk.LEFT, padx=10)

# 生成修改建议的函数
def generate_revision_suggestions(evaluation_result):
    if not generated_content.strip():
        messagebox.showinfo("提示", "没有可以修改的内容")
        return
    
    update_status("正在生成修改建议...")
    
    # 创建一个新线程来执行修改建议生成
    thread = threading.Thread(target=lambda: revision_suggestions_thread(evaluation_result))
    thread.daemon = True
    thread.start()

# 生成修改建议的线程函数
def revision_suggestions_thread(evaluation_result):
    try:
        user_prompt = prompt_entry.get("1.0", tk.END).strip()
        
        # 构造修改建议提示词
        revision_prompt = f'''
        你是一个专业的小说编辑，请基于以下评估报告，对小说内容提供具体的修改建议。
        
        ## 小说评估报告：
        {evaluation_result}
        
        ## 原始写作要求：
        {user_prompt}
        
        ## 小说内容：
        {generated_content}
        
        请针对评估报告中指出的问题，提供以下内容：
        1. 对小说结构的修改建议
        2. 对情节发展的改进方案
        3. 对人物刻画的增强建议
        4. 对语言表达的优化方案
        5. 提供3个具体的修改示例，包括原文和修改后的对比
        
        请确保你的建议是具体且可操作的，而不是笼统的指导。
        '''
        
        # 准备请求参数
        api_url = 'http://localhost:11434/api/generate'
        request_data = {
            "model": evaluation_model_name,
            "prompt": revision_prompt,
            "temperature": 0.4,
            "stream": False
        }
        
        # 发送请求
        response = requests.post(api_url, json=request_data, timeout=600)
        response.raise_for_status()
        
        # 获取修改建议
        result = response.json()
        revision_suggestions = result.get('response', '').strip()
        
        # 显示修改建议
        show_revision_suggestions(revision_suggestions)
        
        update_status("修改建议生成完成")
    except Exception as e:
        messagebox.showerror("错误", f"生成修改建议时出错: {str(e)}")
        update_status("生成修改建议时出错")

# 显示修改建议的函数
def show_revision_suggestions(suggestions):
    # 创建修改建议窗口
    suggestions_window = tk.Toplevel(root)
    suggestions_window.title("小说修改建议")
    suggestions_window.geometry("800x600")
    suggestions_window.configure(bg='#f5f5f7')
    
    # 创建主容器
    suggestions_container = tk.Frame(suggestions_window, bg='#f5f5f7')
    suggestions_container.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)
    
    # 创建标题
    tk.Label(
        suggestions_container,
        text="小说修改建议",
        font=('Microsoft YaHei UI', 16, 'bold'),
        bg='#f5f5f7',
        fg='#333333'
    ).pack(pady=(0, 15))
    
    # 创建修改建议文本区域
    suggestions_text = scrolledtext.ScrolledText(
        suggestions_container,
        wrap=tk.WORD,
        width=80,
        height=25,
        font=('Microsoft YaHei UI', 11),
        bg='#ffffff',
        fg='#333333',
        padx=10,
        pady=10
    )
    suggestions_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
    suggestions_text.insert(tk.END, suggestions)
    
    # 创建按钮区域
    button_frame = tk.Frame(suggestions_container, bg='#f5f5f7')
    button_frame.pack(pady=15)
    
    # 添加应用修改按钮
    tk.Button(
        button_frame,
        text="应用修改",
        command=lambda: apply_revisions(suggestions),
        font=('Microsoft YaHei UI', 10, 'bold'),
        bg='#34c759',
        fg='white',
        relief=tk.RAISED,
        bd=0,
        padx=20,
        pady=8,
        cursor="hand2"
    ).pack(side=tk.LEFT, padx=10)
    
    # 添加关闭按钮
    tk.Button(
        button_frame,
        text="关闭",
        command=suggestions_window.destroy,
        font=('Microsoft YaHei UI', 10, 'bold'),
        bg='#8e8e93',
        fg='white',
        relief=tk.RAISED,
        bd=0,
        padx=20,
        pady=8,
        cursor="hand2"
    ).pack(side=tk.LEFT, padx=10)

# 应用修改的函数
def apply_revisions(suggestions):
    global generated_content, thinking_content
    
    if not generated_content.strip():
        messagebox.showinfo("提示", "没有可以修改的内容")
        return
    
    # 询问用户是否要继续
    if not messagebox.askyesno("确认", "这将基于修改建议重写小说内容，是否继续？"):
        return
    
    update_status("正在应用修改...")
    
    # 创建一个新线程来执行应用修改
    thread = threading.Thread(target=lambda: apply_revisions_thread(suggestions))
    thread.daemon = True
    thread.start()

# 应用修改的线程函数
def apply_revisions_thread(suggestions):
    global generated_content, thinking_content
    
    try:
        user_prompt = prompt_entry.get("1.0", tk.END).strip()
        
        # 构造修改提示词
        revision_prompt = f'''
        你是一个专业的小说编辑和作家，请基于以下修改建议，重写小说内容。
        
        ## 原始写作要求：
        {user_prompt}
        
        ## 原始小说内容：
        {generated_content}
        
        ## 修改建议：
        {suggestions}
        
        请根据以上修改建议，重写整个小说。保留原有的故事框架和主要情节，但根据修改建议进行优化和改进。
        注意：
        1. 请直接输出修改后的完整小说内容，不要包含任何解释或说明
        2. 优化情节、人物和语言表达，但保持故事的连贯性和原意
        3. 保持适当的篇幅，不要过度缩减或扩展内容
        4. 使用流畅的中文写作
        5. 你可以使用<think></think>标签来表示你的思考过程，这部分内容不会出现在最终故事中。
        '''
        
        # 准备请求参数
        api_url = 'http://localhost:11434/api/generate'
        request_data = {
            "model": writing_model_name,
            "prompt": revision_prompt,
            "temperature": 0.5,
            "stream": True
        }
        
        # 清空显示区域
        output_text.delete(1.0, tk.END)
        thinking_text.delete(1.0, tk.END)
        revised_content = ""
        thinking_content = ""
        
        # 发送请求并获取流式响应
        response = requests.post(api_url, json=request_data, stream=True, timeout=120)
        response.raise_for_status()
        
        # 处理流式响应
        for line in response.iter_lines():
            if line:
                # 解析 JSON 数据
                data = json.loads(line.decode('utf-8'))
                # 提取 response 字段
                text_chunk = data.get('response', '')
                
                # 处理文本块，分离思维推理和正式内容
                story_chunk = process_text_chunk(text_chunk)
                
                # 更新显示区域（仅显示正式内容）
                if story_chunk:
                    output_text.insert(tk.END, story_chunk)
                    output_text.see(tk.END)  # 自动滚动到最新内容
                    # 累积新内容
                    revised_content += story_chunk
                
                root.update_idletasks()  # 更新UI
                
                # 检查是否完成
                if data.get('done', False):
                    break
        
        # 更新生成的内容
        generated_content = revised_content
        
        # 保存修改后的内容
        save_content_to_file()
        
        update_status("小说内容已修改并保存")
    except Exception as e:
        messagebox.showerror("错误", f"应用修改时出错: {str(e)}")
        update_status("应用修改时出错")

# 修改update_status函数以支持更多状态样式
def update_status(message):
    status_label.config(text=message)
    if "错误" in message:
        status_label.config(bg='#ffebee', fg='#c62828')  # 错误状态使用红色
    elif "完成" in message:
        status_label.config(bg='#e8f5e9', fg='#2e7d32')  # 完成状态使用绿色
    elif "生成中" in message:
        status_label.config(bg='#e3f2fd', fg='#1565c0')  # 生成中状态使用蓝色
    elif "保存" in message:
        status_label.config(bg='#fff8e1', fg='#f57f17')  # 保存状态使用琥珀色
    else:
        status_label.config(bg='#e8e8e8', fg='#333333')  # 默认状态
    root.update_idletasks()

# 复制内容到剪贴板
def copy_content_to_clipboard():
    if not generated_content.strip():
        messagebox.showinfo("提示", "没有可复制的内容")
        return
    
    try:
        pyperclip.copy(generated_content)
        update_status("内容已复制到剪贴板")
    except Exception as e:
        messagebox.showerror("错误", f"复制内容时出错: {str(e)}")
        update_status("复制内容时出错")

# 切换Markdown模式
def toggle_markdown_mode():
    global is_markdown_mode, output_text
    
    if not generated_content.strip():
        messagebox.showinfo("提示", "没有内容可以进行格式转换")
        return
    
    is_markdown_mode = not is_markdown_mode
    
    try:
        # 保存当前滚动位置
        current_scroll_position = output_text.yview()[0]
        
        # 从PanedWindow中移除当前的文本控件
        paned_window.forget(output_text)
        
        if is_markdown_mode:
            # 切换到Markdown模式
            markdown_button.config(text="普通模式阅读")
            
            # 创建HTML显示控件
            html_content = markdown.markdown(generated_content)
            output_text = HTMLScrolledText(paned_window, html=html_content, padx=5, pady=5)
            paned_window.add(output_text, stretch="always", width=600)
            
            update_status("已切换到Markdown格式显示")
        else:
            # 切换回普通模式
            markdown_button.config(text="Markdown格式阅读")
            
            # 创建普通文本控件
            output_text = scrolledtext.ScrolledText(
                paned_window,
                wrap=tk.WORD,
                width=70,
                height=30,
                font=('Microsoft YaHei UI', 11),
                bg='#f9f9f9',
                fg='#333333',
                padx=3,
                pady=3,
                relief=tk.FLAT
            )
            paned_window.add(output_text, stretch="always", width=600)
            
            # 填充内容
            output_text.insert(tk.END, generated_content)
            
            # 配置滚动条
            scrollbar1 = ttk.Scrollbar(output_text)
            output_text.config(yscrollcommand=scrollbar1.set)
            scrollbar1.config(command=output_text.yview)
            
            update_status("已切换到普通模式显示")
        
        # 尝试恢复滚动位置
        try:
            output_text.yview_moveto(current_scroll_position)
        except:
            pass
            
    except Exception as e:
        messagebox.showerror("错误", f"切换显示模式时出错: {str(e)}")
        update_status("切换显示模式时出错")

# 导出为DOCX文件
def export_to_docx():
    if not generated_content.strip():
        messagebox.showinfo("提示", "没有可导出的内容")
        return
    
    try:
        # 弹出文件保存对话框
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx")],
            title="保存为Word文档"
        )
        
        if not file_path:  # 用户取消了操作
            return
        
        # 创建一个新的Word文档
        doc = docx.Document()
        
        # 添加标题
        title = prompt_entry.get("1.0", "end-1c")[:50] + "..."  # 截取提示的前50个字符作为标题
        doc.add_heading(title, level=1)
        
        # 设置正文样式
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        
        # 分段添加内容
        paragraphs = generated_content.split('\n\n')
        for para in paragraphs:
            if para.strip():  # 确保段落不是空的
                doc.add_paragraph(para)
        
        # 保存文档
        doc.save(file_path)
        
        update_status(f"已成功导出为Word文档: {os.path.basename(file_path)}")
    except Exception as e:
        messagebox.showerror("错误", f"导出Word文档时出错: {str(e)}")
        update_status("导出Word文档时出错")

# 创建主窗口
root = tk.Tk()
root.title("AI 长篇小说写作助手")
root.geometry("1000x800")
root.configure(bg='#f5f5f7')  # 更新为浅灰背景色

# 设置样式
style = ttk.Style()
style.configure('TFrame', background='#f5f5f7')
style.configure('TLabel', background='#f5f5f7', font=('Microsoft YaHei UI', 10))
style.configure('TButton', font=('Microsoft YaHei UI', 10), padding=6)
style.configure('Primary.TButton', background='#007aff', foreground='white')
style.configure('Danger.TButton', background='#ff3b30', foreground='white')
style.configure('Success.TButton', background='#34c759', foreground='white')

# 创建主容器
main_container = tk.Frame(root, bg='#f5f5f7')
main_container.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)

# 创建标题区域，带有装饰线条
title_frame = tk.Frame(main_container, bg='#f5f5f7')
title_frame.pack(fill=tk.X, pady=(0, 8)) # 添加间距 

title_label = tk.Label(
    title_frame,
    text="AI 长篇小说写作助手",
    font=('Microsoft YaHei UI', 10, 'bold'),        
    bg='#f5f5f7',
    fg='#333333'
)
title_label.pack(pady=(0, 8))

# 添加分隔线
separator = ttk.Separator(title_frame, orient='horizontal')
separator.pack(fill=tk.X, pady=(0, 8))

# 创建输入区域框架，使用圆角和阴影效果
input_frame = tk.LabelFrame(
    main_container,
    text="写作设置",
    font=('Microsoft YaHei UI', 11, 'bold'),
    bg='#ffffff',
    fg='#333333',
    padx=5,
    pady=5,
    relief=tk.GROOVE,
    bd=2
)
input_frame.pack(fill=tk.X, padx=5, pady=(0, 8))

# 创建提示词输入区域
prompt_label = tk.Label(
    input_frame,
    text="写作要求：",
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#ffffff',
    fg='#333333'
)
prompt_label.pack(side=tk.LEFT, padx=(5, 10), anchor='n')

# 创建一个Frame来容纳文本框
prompt_container = tk.Frame(input_frame, bg='#ffffff')
prompt_container.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)

prompt_entry = tk.Text(
    prompt_container,
    width=50,
    height=5,  # 增加高度为更好的可读性
    font=('Microsoft YaHei UI', 10),
    wrap=tk.WORD,
    relief=tk.SOLID,
    bd=1,
    padx=5,
    pady=5,
    bg='#f9f9f9'
)
prompt_entry.pack(fill=tk.BOTH, expand=True)

# 创建右侧设置区域
settings_frame = tk.Frame(input_frame, bg='#ffffff', padx=10)
settings_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(5, 5))

# 创建字数设置区域
word_count_label = tk.Label(
    settings_frame,
    text="目标字数：",
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#ffffff',
    fg='#333333'
)
word_count_label.pack(side=tk.TOP, anchor='w', pady=(0, 8))

word_count_entry = tk.Entry(
    settings_frame,
    width=15,
    font=('Microsoft YaHei UI', 10),
    relief=tk.SOLID,
    bd=1,
    bg='#f9f9f9'
)
word_count_entry.pack(side=tk.TOP, anchor='w', pady=(0, 8))
word_count_entry.insert(0, "3000")

# 创建控制按钮区域，使用卡片式设计
button_card = tk.Frame(main_container, bg='#ffffff', relief=tk.RAISED, bd=1)
button_card.pack(fill=tk.X, padx=5, pady=5)

button_frame = tk.Frame(button_card, bg='#ffffff', padx=5, pady=5)
button_frame.pack(fill=tk.X)

# 添加生成按钮
generate_button = tk.Button(
    button_frame,
    text="开始生成",
    command=generate_text,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#007aff',
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
generate_button.pack(side=tk.LEFT, padx=5)

# 添加停止按钮
stop_button = tk.Button(
    button_frame,
    text="停止生成",
    command=stop_generation,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#ff3b30',
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
stop_button.pack(side=tk.LEFT, padx=5)

# 添加自动生成按钮
auto_generate_button = tk.Button(
    button_frame,
    text="自动生成",
    command=toggle_auto_generate,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#34c759',
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
auto_generate_button.pack(side=tk.LEFT, padx=5)

# 添加复制按钮
copy_button = tk.Button(
    button_frame,
    text="复制内容",
    command=copy_content_to_clipboard,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#5856d6',  # 紫蓝色背景
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
copy_button.pack(side=tk.LEFT, padx=5)

# 添加Markdown切换按钮
markdown_button = tk.Button(
    button_frame,
    text="Markdown格式阅读",
    command=toggle_markdown_mode,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#ff9500',  # 橙色背景
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
markdown_button.pack(side=tk.LEFT, padx=5)

# 添加导出DOCX按钮
export_docx_button = tk.Button(
    button_frame,
    text="导出DOCX",
    command=export_to_docx,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#0098f7',  # 蓝色背景
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
export_docx_button.pack(side=tk.LEFT, padx=5)

# 添加质量评估按钮
evaluation_button = tk.Button(
    button_frame,
    text="质量评估",
    command=evaluate_novel_quality,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#9C27B0',  # 紫色背景
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
evaluation_button.pack(side=tk.LEFT, padx=5)

# 添加退出按钮
exit_button = tk.Button(
    button_frame,
    text="退出程序",
    command=root.destroy,
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#8e8e93',
    fg='white',
    relief=tk.RAISED,
    bd=0,
    padx=5,
    pady=5,
    cursor="hand2"
)
exit_button.pack(side=tk.RIGHT, padx=5)

# 添加状态标签（放在按钮右边）
status_label = tk.Label(
    button_frame,
    text="就绪",
    font=('Microsoft YaHei UI', 10),
    bg='#e8e8e8',
    fg='#333333',
    relief=tk.GROOVE,
    padx=5,
    pady=5
)
status_label.pack(side=tk.RIGHT, padx=5, fill=tk.X, expand=True)

# 创建显示区域，使用卡片式设计
output_frame = tk.LabelFrame(
    main_container,
    text="写作内容",
    font=('Microsoft YaHei UI', 11, 'bold'),
    bg='#ffffff',
    fg='#333333',
    padx=3,
    pady=3,
    relief=tk.GROOVE,
    bd=2
)
output_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

# 创建一个PanedWindow来分割小说内容和思维推理内容
paned_window = tk.PanedWindow(output_frame, orient=tk.HORIZONTAL, bg='#ffffff', sashwidth=4, sashrelief=tk.RAISED)
paned_window.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

# 创建小说内容文本区域
output_text = scrolledtext.ScrolledText(
    paned_window,
    wrap=tk.WORD,
    width=70,
    height=30,
    font=('Microsoft YaHei UI', 11),
    bg='#f9f9f9',
    fg='#333333',
    padx=3,
    pady=3,
    relief=tk.FLAT
)
paned_window.add(output_text, stretch="always", width=600)

# 创建思维推理内容框架
thinking_frame = tk.Frame(paned_window, bg='#ffffff')
thinking_frame.pack(fill=tk.BOTH, expand=True)
paned_window.add(thinking_frame, width=300)

# 思维推理标题
tk.Label(
    thinking_frame,
    text="思维推理过程",
    font=('Microsoft YaHei UI', 10, 'bold'),
    bg='#ffffff',
    fg='#333333'
).pack(fill=tk.X, pady=(0, 5))

# 创建思维推理内容文本区域
thinking_text = scrolledtext.ScrolledText(
    thinking_frame,
    wrap=tk.WORD,
    width=30,
    height=30,
    font=('Microsoft YaHei UI', 10),
    bg='#f0f0f0',
    fg='#555555',
    padx=3,
    pady=3,
    relief=tk.FLAT
)
thinking_text.pack(fill=tk.BOTH, expand=True)

# 自定义滚动条样式
scrollbar1 = ttk.Scrollbar(output_text)
output_text.config(yscrollcommand=scrollbar1.set)
scrollbar1.config(command=output_text.yview)

scrollbar2 = ttk.Scrollbar(thinking_text)
thinking_text.config(yscrollcommand=scrollbar2.set)
scrollbar2.config(command=thinking_text.yview)

# 启动主循环
root.mainloop() 